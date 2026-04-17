# src/processing/extract_docx.py
from __future__ import annotations

import io
import re
from typing import Dict, List, Tuple

from docx import Document
from docx.oxml.ns import qn

# ── constants ──────────────────────────────────────────────────────────────────

# Word built-in heading style names
_HEADING_STYLES = {
    "heading 1", "heading 2", "heading 3",
    "heading 4", "heading 5", "heading 6",
    "title", "subtitle",
}

# Caption style names
_CAPTION_STYLES = {"caption"}

# Figure/table caption heuristic (same as PDF)
_CAPTION_RE = re.compile(
    r"^\s*(?:Fig\.?|Figure|TABLE|Table)\s*[\dIVX]+[\.\:]",
    re.IGNORECASE,
)

# IEEE-style section heading in plain text
_SECTION_RE = re.compile(
    r"^\s*(?:"
    r"(?:[IVX]{1,6}\.)\s+[A-Z]"
    r"|(?:[A-Z]\.)\s+[A-Z]"
    r"|(?:\d{1,2}\.\d{0,2})\s+\S"
    r")\s*",
    re.IGNORECASE,
)


# ── helpers ────────────────────────────────────────────────────────────────────

def _para_style(para) -> str:
    """Return normalised paragraph style name."""
    try:
        return (para.style.name or "").strip().lower()
    except Exception:
        return ""


def _is_heading(para) -> bool:
    style = _para_style(para)
    if style in _HEADING_STYLES:
        return True
    # fallback: outline level in XML
    try:
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            outLvl = pPr.find(qn("w:outlineLvl"))
            if outLvl is not None:
                lvl = int(outLvl.get(qn("w:val"), 9))
                return lvl < 6
    except Exception:
        pass
    # fallback: text pattern
    t = (para.text or "").strip()
    if _SECTION_RE.match(t):
        return True
    return False


def _is_caption(para) -> bool:
    style = _para_style(para)
    if style in _CAPTION_STYLES:
        return True
    t = (para.text or "").strip()
    return bool(_CAPTION_RE.match(t))


def _table_to_markdown(table) -> str:
    """Convert a docx table to a markdown string."""
    rows = []
    for row in table.rows:
        cells = [
            (cell.text or "").strip().replace("\n", " ")
            for cell in row.cells
        ]
        rows.append(cells)

    if not rows:
        return ""

    # deduplicate merged cells (docx repeats merged cell text)
    clean_rows = []
    for row in rows:
        seen = set()
        clean = []
        for i, c in enumerate(row):
            if i > 0 and c == row[i - 1] and c in seen:
                clean.append("")
            else:
                clean.append(c)
                seen.add(c)
        clean_rows.append(clean)

    max_cols = max(len(r) for r in clean_rows)
    # pad rows
    clean_rows = [(r + [""] * max_cols)[:max_cols] for r in clean_rows]

    header = clean_rows[0]
    sep    = ["---"] * max_cols
    lines  = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep)    + " |",
    ]
    for row in clean_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _to_bytesio(data: bytes) -> io.BytesIO:
    return io.BytesIO(data)


# ── public API (Stage 3 contract) ─────────────────────────────────────────────

def extract_docx_pages(docx_bytes: bytes) -> Tuple[List[Dict], str]:
    """
    Stage 3 contract (unchanged signature):
      returns (pages, full_text_original)

    DOCX has no reliable page boundaries → we produce logical sections
    as separate 'pages' when headings are found, otherwise one page.

    Each page dict:
      page_num, text_original, text_clean="", flags={}
      + NEW optional fields: figures=[], tables=[], sections=[]

    runner.py and chunk.py need zero changes.
    """
    doc = Document(_to_bytesio(docx_bytes))

    # ── collect document children in order ────────────────────────────────────
    # We interleave paragraphs and tables in document order using the XML body
    from docx.oxml.ns import qn as _qn
    body = doc.element.body

    # Build index maps so we can look up by XML element
    para_map  = {p._p: p  for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}

    # ── pass 1: collect ordered items ─────────────────────────────────────────
    ordered_items = []  # each: {"kind": "para"|"table", "obj": ...}
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p" and child in para_map:
            ordered_items.append({"kind": "para", "obj": para_map[child]})
        elif tag == "tbl" and child in table_map:
            ordered_items.append({"kind": "table", "obj": table_map[child]})

    # ── pass 2: group into logical sections (split at headings) ───────────────
    # Each section becomes one "page" for the pipeline
    sections_raw: List[List] = [[]]   # list of item groups

    for item in ordered_items:
        if item["kind"] == "para" and _is_heading(item["obj"]):
            t = (item["obj"].text or "").strip()
            if t:
                # start new section only if current section has content
                if any(
                    i["kind"] == "table"
                    or (i["kind"] == "para" and (i["obj"].text or "").strip())
                    for i in sections_raw[-1]
                ):
                    sections_raw.append([])
        sections_raw[-1].append(item)

    # ── pass 3: render each section into a page dict ──────────────────────────
    pages:      List[Dict] = []
    all_parts:  List[str]  = []
    page_num = 0

    for section_items in sections_raw:
        parts:    List[str]  = []
        figures:  List[Dict] = []
        tables:   List[Dict] = []
        sec_headings: List[str] = []

        for item in section_items:
            if item["kind"] == "para":
                para = item["obj"]
                t    = (para.text or "").strip()
                if not t:
                    continue

                if _is_heading(para):
                    sec_headings.append(t)
                    parts.append(t)

                elif _is_caption(para):
                    figures.append({"caption": t, "page_num": page_num + 1})
                    parts.append(t)

                else:
                    parts.append(t)

            elif item["kind"] == "table":
                md = _table_to_markdown(item["obj"])
                if md:
                    tables.append({
                        "markdown":  md,
                        "row_count": len(item["obj"].rows),
                        "col_count": len(item["obj"].columns),
                    })
                    parts.append("\n" + md + "\n")

        text_original = "\n\n".join(p for p in parts if p.strip()).strip()
        if not text_original:
            continue

        page_num += 1
        all_parts.append(text_original)

        flags = {
            "method":         "docx_structured",
            "warnings":       [],
            "has_figures":    len(figures) > 0,
            "has_tables":     len(tables)  > 0,
            "has_formulas":   False,
            "sections_found": sec_headings,
        }

        pages.append({
            "page_num":      page_num,
            "text_original": text_original,
            "text_clean":    "",
            "flags":         flags,
            "figures":       figures,
            "tables":        tables,
            "sections":      sec_headings,
        })

    # ── fallback: no sections found → single page ─────────────────────────────
    if not pages:
        full = "\n\n".join(all_parts).strip()
        pages = [{
            "page_num":      1,
            "text_original": full,
            "text_clean":    "",
            "flags":         {"method": "docx_flat", "warnings": []},
            "figures":       [],
            "tables":        [],
            "sections":      [],
        }]
        all_parts = [full]

    full_text_original = "\n\n".join(t for t in all_parts if t).strip()
    return pages, full_text_original